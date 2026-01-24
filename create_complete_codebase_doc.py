from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import glob

def add_heading_styled(doc, text, level, color=(25, 55, 109)):
    """Add styled heading"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(*color)
        run.font.name = 'Segoe UI'
    heading.space_before = Pt(12)
    heading.space_after = Pt(6)
    return heading

def add_code_block(doc, code_text, language="python"):
    """Add a code block with syntax highlighting style"""
    p = doc.add_paragraph()
    p.style = 'Normal'
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(8)  # Smaller font for code
    run.font.color.rgb = RGBColor(0, 100, 0)  # Dark green for code
    return p

def add_file_content(doc, file_path, relative_path):
    """Add a file's content to the document"""
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()

        # Add file header
        file_header = doc.add_paragraph()
        file_header_run = file_header.add_run(f"📄 {relative_path}")
        file_header_run.font.name = 'Segoe UI'
        file_header_run.font.size = Pt(12)
        file_header_run.font.bold = True
        file_header_run.font.color.rgb = RGBColor(25, 55, 109)
        file_header.space_before = Pt(18)
        file_header.space_after = Pt(6)

        # Add file size info
        file_size = len(content)
        size_info = doc.add_paragraph()
        size_run = size_info.add_run(f"File size: {file_size:,} characters")
        size_run.font.name = 'Segoe UI'
        size_run.font.size = Pt(9)
        size_run.font.color.rgb = RGBColor(100, 100, 100)
        size_info.space_after = Pt(6)

        # Add content
        if content.strip():
            add_code_block(doc, content)
        else:
            empty_note = doc.add_paragraph()
            empty_run = empty_note.add_run("(Empty file)")
            empty_run.font.name = 'Segoe UI'
            empty_run.font.size = Pt(9)
            empty_run.font.italic = True
            empty_run.font.color.rgb = RGBColor(150, 150, 150)

        # Add separator
        separator = doc.add_paragraph()
        sep_run = separator.add_run("=" * 80)
        sep_run.font.name = 'Consolas'
        sep_run.font.size = Pt(8)
        sep_run.font.color.rgb = RGBColor(200, 200, 200)
        separator.space_before = Pt(12)
        separator.space_after = Pt(12)

    except Exception as e:
        error_para = doc.add_paragraph()
        error_run = error_para.add_run(f"Error reading {relative_path}: {str(e)}")
        error_run.font.name = 'Segoe UI'
        error_run.font.size = Pt(9)
        error_run.font.color.rgb = RGBColor(200, 0, 0)

def get_code_files(base_path):
    """Get all code files to include in the document"""
    code_files = []

    # Python files
    code_files.extend(glob.glob(os.path.join(base_path, '**', '*.py'), recursive=True))

    # TypeScript/JavaScript files
    code_files.extend(glob.glob(os.path.join(base_path, '**', '*.ts'), recursive=True))
    code_files.extend(glob.glob(os.path.join(base_path, '**', '*.tsx'), recursive=True))
    code_files.extend(glob.glob(os.path.join(base_path, '**', '*.js'), recursive=True))
    code_files.extend(glob.glob(os.path.join(base_path, '**', '*.jsx'), recursive=True))

    # Configuration files
    code_files.extend(glob.glob(os.path.join(base_path, '**', '*.json'), recursive=True))
    code_files.extend(glob.glob(os.path.join(base_path, '**', '*.yaml'), recursive=True))
    code_files.extend(glob.glob(os.path.join(base_path, '**', '*.yml'), recursive=True))
    code_files.extend(glob.glob(os.path.join(base_path, '**', '*.toml'), recursive=True))
    code_files.extend(glob.glob(os.path.join(base_path, '**', '*.ini'), recursive=True))
    code_files.extend(glob.glob(os.path.join(base_path, '**', '*.cfg'), recursive=True))
    code_files.extend(glob.glob(os.path.join(base_path, '**', '*.conf'), recursive=True))

    # Infrastructure files
    code_files.extend(glob.glob(os.path.join(base_path, '**', '*.bicep'), recursive=True))
    code_files.extend(glob.glob(os.path.join(base_path, '**', '*.tf'), recursive=True))

    # Shell scripts
    code_files.extend(glob.glob(os.path.join(base_path, '**', '*.sh'), recursive=True))
    code_files.extend(glob.glob(os.path.join(base_path, '**', '*.ps1'), recursive=True))
    code_files.extend(glob.glob(os.path.join(base_path, '**', '*.bat'), recursive=True))

    # Documentation files
    code_files.extend(glob.glob(os.path.join(base_path, '**', '*.md'), recursive=True))

    # Exclude certain files/directories
    exclude_patterns = [
        '__pycache__',
        'node_modules',
        '.git',
        '.venv',
        'build',
        'dist',
        '*.pyc',
        '*.pyo',
        '*.log',
        '*.tmp',
        '*.swp',
        '*.bak',
        '*.orig'
    ]

    filtered_files = []
    for file_path in code_files:
        should_exclude = False
        for pattern in exclude_patterns:
            if pattern in file_path:
                should_exclude = True
                break
        if not should_exclude:
            filtered_files.append(file_path)

    return sorted(filtered_files)

# ==================== WORD - COMPLETE CODEBASE ====================
doc = Document()
doc.default_tab_stop = Inches(0.5)

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Segoe UI'
font.size = Pt(10)

# Title
title = doc.add_paragraph()
title_run = title.add_run('KRAFTD DOCS')
title_run.font.size = Pt(32)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(25, 55, 109)
title_run.font.name = 'Segoe UI'
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
title.space_after = Pt(2)

subtitle = doc.add_paragraph()
subtitle_run = subtitle.add_run('Complete Source Code Repository')
subtitle_run.font.size = Pt(18)
subtitle_run.font.italic = True
subtitle_run.font.color.rgb = RGBColor(100, 100, 100)
subtitle_run.font.name = 'Segoe UI'
subtitle.space_after = Pt(8)

date_line = doc.add_paragraph()
date_run = date_line.add_run('January 23, 2026')
date_run.font.size = Pt(10)
date_run.font.color.rgb = RGBColor(150, 150, 150)
date_run.font.name = 'Segoe UI'

# Introduction
intro = doc.add_paragraph()
intro_run = intro.add_run(
    'This document contains the complete source code of the Kraftd Docs SaaS platform, '
    'organized by directory structure. All Python, TypeScript, JavaScript, configuration, '
    'infrastructure, and documentation files are included.'
)
intro_run.font.name = 'Segoe UI'
intro_run.font.size = Pt(11)

# Get base path
base_path = r'c:\Users\1R6\OneDrive\Project Catalyst\KraftdIntel'

# Get all code files
code_files = get_code_files(base_path)

# Organize by directory
file_structure = {}
for file_path in code_files:
    relative_path = os.path.relpath(file_path, base_path)
    directory = os.path.dirname(relative_path)
    if directory not in file_structure:
        file_structure[directory] = []
    file_structure[directory].append((file_path, relative_path))

# Process each directory
for directory in sorted(file_structure.keys()):
    if directory == '':
        dir_name = 'Root Directory'
    else:
        dir_name = directory.replace('\\', ' / ')

    add_heading_styled(doc, f'📁 {dir_name}', 1, (25, 55, 109))

    # Add files in this directory
    for file_path, relative_path in sorted(file_structure[directory]):
        add_file_content(doc, file_path, relative_path)

# Summary
add_heading_styled(doc, '📊 Summary', 1)

summary_text = doc.add_paragraph()
summary_run = summary_text.add_run(
    f'Total files included: {len(code_files)}\n'
    f'Directories covered: {len(file_structure)}\n'
    f'Generated on: January 23, 2026'
)
summary_run.font.name = 'Segoe UI'
summary_run.font.size = Pt(11)

# Save the document
output_path = r'c:\Users\1R6\OneDrive\Desktop\Kraftd Docs\KRAFTD_Complete_Source_Code.docx'
doc.save(output_path)

print(f"Complete source code document created successfully at: {output_path}")
print(f"Total files included: {len(code_files)}")
print("Opening document...")

# Open the document
os.startfile(output_path)