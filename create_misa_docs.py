from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_heading_styled(doc, text, level, color=(25, 55, 109)):
    """Add styled heading"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(*color)
        run.font.name = 'Segoe UI'
    heading.space_before = Pt(12)
    heading.space_after = Pt(6)
    return heading

def set_cell_background(cell, fill_color):
    """Set cell background color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill_color)
    cell._element.get_or_add_tcPr().append(shading_elm)

# ==================== WORD - COMPANY PROFILE ====================
doc_profile = Document()
doc_profile.default_tab_stop = Inches(0.5)

# Set default font
style = doc_profile.styles['Normal']
font = style.font
font.name = 'Segoe UI'
font.size = Pt(11)

# Title
title = doc_profile.add_paragraph()
title_run = title.add_run('KRAFTD')
title_run.font.size = Pt(32)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(25, 55, 109)
title_run.font.name = 'Segoe UI'
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
title.space_after = Pt(2)

subtitle = doc_profile.add_paragraph()
subtitle_run = subtitle.add_run('Intelligent Supply Chain for SMEs')
subtitle_run.font.size = Pt(14)
subtitle_run.font.italic = True
subtitle_run.font.color.rgb = RGBColor(100, 100, 100)
subtitle_run.font.name = 'Segoe UI'
subtitle.space_after = Pt(8)

divider = doc_profile.add_paragraph('_' * 70)
divider.space_after = Pt(12)

# Meta info
meta_para = doc_profile.add_paragraph()
meta_para.add_run('MISA Entrepreneurship License Application  |  January 2026  |  Saudi Arabia').font.size = Pt(10)
meta_para.space_after = Pt(16)

# THE FOUNDING INSIGHT
add_heading_styled(doc_profile, 'The Founding Insight', 1, (25, 55, 109))
insight = doc_profile.add_paragraph('''SMEs buy raw materials, process them, export them, and deliver goods—exactly like SABIC. Yet the proportion of intelligence available to them is nowhere close.

SABIC's supply chain data is structured, connected, and accessible through high-tech endpoints. SMEs performing identical activities operate without equivalent data visibility or infrastructure.

This is the inequality Kraftd solves.''')
for run in insight.runs:
    run.font.name = 'Segoe UI'
    run.font.size = Pt(11)
insight.space_after = Pt(12)

# COMPANY IDENTITY
add_heading_styled(doc_profile, 'What Kraftd Does', 1, (25, 55, 109))
doc_profile.add_paragraph('Transform unstructured business documents (invoices, quotations, BOQs, contracts) into structured, actionable intelligence.', style='List Bullet')
doc_profile.add_paragraph('Provide organizations oversight into how AI tools operate in workflows—ensuring safety and reliability.', style='List Bullet')
doc_profile.add_paragraph('Enable SMEs to operate with the clarity and precision previously available only to large enterprises.', style='List Bullet')

# FOUNDER
add_heading_styled(doc_profile, 'The Founder', 1, (25, 55, 109))
doc_profile.add_paragraph('Sulaiman Ul Akram Ismail – 12+ years in petrochemical, logistics, and industrial operations across the GCC', style='List Bullet')
doc_profile.add_paragraph('✓ Worked with Saudi Aramco, SABIC, NEOM | Zero-incident safety record (500,000+ safe man-hours)', style='List Bullet')
doc_profile.add_paragraph('✓ Mechanical Engineering degree | McKinsey Forward Program | Founder Institute Launch Track', style='List Bullet')
doc_profile.add_paragraph('✓ Built entire MVP single-handedly (full-stack development)', style='List Bullet')
doc_profile.add_paragraph('✓ Lived the problem for a decade – knows SME pain intimately', style='List Bullet')

# THE PROBLEM
add_heading_styled(doc_profile, 'The Problem: Why Digitalization Failed SMEs', 1, (25, 55, 109))
problem_para = doc_profile.add_paragraph('''When digitalization was imposed on SMEs—portals, uploads, templates, compliance tools—it didn't make them efficient. It added more work, more cost, and more administrative burden just to meet client requirements.

The reason: AI and ML platforms built ecosystems around APIs because it's cheap and scalable. But this intelligence only serves companies that already have structured systems. It never reaches SMEs.

SAP, Oracle, and vendor platforms tailor data flows to benefit large enterprises, not the suppliers feeding them. These systems were never constructed the other way around.

Result: SMEs spend 15-30% of operational time on document processing, lack visibility into decision quality, and cannot scale efficiently without enterprise-grade infrastructure.''')
for run in problem_para.runs:
    run.font.name = 'Segoe UI'
    run.font.size = Pt(11)
problem_para.space_after = Pt(12)

# MARKET OPPORTUNITY
add_heading_styled(doc_profile, 'Market Opportunity', 1, (25, 55, 109))
table = doc_profile.add_table(rows=5, cols=3)
table.style = 'Light Grid Accent 1'
header_cells = table.rows[0].cells
header_cells[0].text = 'Metric'
header_cells[1].text = 'Value'
header_cells[2].text = 'Implication'
for cell in header_cells:
    set_cell_background(cell, 'E8EEF7')
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.size = Pt(10)

data = [
    ['SME Population in Saudi Arabia', '600,000+', 'Addressable market scale'],
    ['Annual Procurement Value', '$400B+', 'Size of opportunity'],
    ['Manual Procurement Process', '78%', 'Digitalization gap'],
    ['Time Wasted on Document Processing', '15-30%', 'Efficiency opportunity']
]
for i, row_data in enumerate(data, 1):
    row_cells = table.rows[i].cells
    for j, cell_text in enumerate(row_data):
        row_cells[j].text = cell_text
        for paragraph in row_cells[j].paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10)
                run.font.name = 'Segoe UI'

doc_profile.add_paragraph()

# THE SOLUTION
add_heading_styled(doc_profile, 'The Solution', 1, (25, 55, 109))
solution = doc_profile.add_paragraph('Kraftd is a cloud-based SaaS platform that works where SMEs actually are—before ERP, before schemas, before infrastructure.')
for run in solution.runs:
    run.font.name = 'Segoe UI'
solution.space_after = Pt(8)
doc_profile.add_paragraph('Document Intelligence Engine: 95%+ accuracy extraction from unstructured documents', style='List Bullet')
doc_profile.add_paragraph('AI Oversight System: Monitor LLM reliability in organizational workflows', style='List Bullet')
doc_profile.add_paragraph('SME-Optimized: 2-week implementation, no IT team required, $500-5K/month cost', style='List Bullet')
doc_profile.add_paragraph('vs. SAP/Oracle: 6-18 months, $50K+/month, requires enterprise infrastructure', style='List Bullet')

doc_profile.add_paragraph()

# TRACTION
add_heading_styled(doc_profile, 'Traction & Validation', 1, (25, 55, 109))
doc_profile.add_paragraph('MVP Live: Production deployment January 2026', style='List Bullet')
doc_profile.add_paragraph('47+ organizations expressing interest | 20+ active trials', style='List Bullet')
doc_profile.add_paragraph('Accuracy: 95%+ on real-world documents', style='List Bullet')
doc_profile.add_paragraph('Customer Willingness to Pay: $3K-15K/month confirmed', style='List Bullet')
doc_profile.add_paragraph('Unit Economics: LTV:CAC = 13:1 to 25:1 | Payback = 2-4 months', style='List Bullet')

doc_profile.add_paragraph()

# FINANCIAL SUMMARY
add_heading_styled(doc_profile, 'Financial Projections', 1, (25, 55, 109))
fin_table = doc_profile.add_table(rows=5, cols=5)
fin_table.style = 'Light Grid Accent 1'
fin_header = fin_table.rows[0].cells
fin_headers = ['Metric', 'Year 1', 'Year 2', 'Year 3', 'Year 5']
for i, header_text in enumerate(fin_headers):
    fin_header[i].text = header_text
    set_cell_background(fin_header[i], 'E8EEF7')
    for paragraph in fin_header[i].paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.size = Pt(10)

fin_data = [
    ['Customers', '15-25', '50-80', '150-250', '500-1K'],
    ['Revenue', '$180-300K', '$600K-1.2M', '$1.8-3M', '$6-12M'],
    ['Profitability', 'Break-even', '+$200-400K EBITDA', '+$1M EBITDA', '+40% EBITDA'],
    ['CAC / LTV', '$400-600 / $8-15K', '$300-500 / $15-25K', '$200-400 / $25-40K', 'Optimized']
]
for i, row_data in enumerate(fin_data, 1):
    row_cells = fin_table.rows[i].cells
    for j, cell_text in enumerate(row_data):
        row_cells[j].text = cell_text
        for paragraph in row_cells[j].paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)
                run.font.name = 'Segoe UI'

doc_profile.add_paragraph()

# CLOSING
add_heading_styled(doc_profile, 'The Opportunity', 1, (25, 55, 109))
closing_text = doc_profile.add_paragraph('''Kraftd solves a structural problem in the supply chain ecosystem: the intelligence inequality between large enterprises and SMEs. This isn't incremental. This is category creation.

With MISA's support, Kraftd becomes the proof point that Saudi Arabia builds world-class technology for global markets. The founder has expertise, the product works, the market is ready.''')
for run in closing_text.runs:
    run.font.name = 'Segoe UI'
    run.font.size = Pt(11)

doc_profile.add_paragraph()
doc_profile.add_paragraph('Contact: akram@kraftd.io | https://kraftd-a4gfhqa2axb2h6cd.uaenorth-01.azurewebsites.net')

doc_profile.save(r'c:\Users\1R6\OneDrive\Desktop\Kraftd Docs\KRAFTD_Company_Profile_MISA.docx')
print("✓ Company Profile (Word) created")

# ==================== WORD - ONE PAGE SUMMARY ====================
doc_summary = Document()
style = doc_summary.styles['Normal']
font = style.font
font.name = 'Segoe UI'
font.size = Pt(10)

title = doc_summary.add_paragraph()
title_run = title.add_run('KRAFTD')
title_run.font.size = Pt(28)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(25, 55, 109)
title.space_after = Pt(2)

subtitle = doc_summary.add_paragraph('Intelligent Supply Chain for SMEs | MISA Application Summary')
for run in subtitle.runs:
    run.font.size = Pt(11)
    run.font.italic = True
    run.font.color.rgb = RGBColor(100, 100, 100)
subtitle.space_after = Pt(10)

# The Problem
heading = doc_summary.add_heading('The Intelligence Inequality', level=2)
for run in heading.runs:
    run.font.color.rgb = RGBColor(25, 55, 109)
    run.font.name = 'Segoe UI'

doc_summary.add_paragraph('''SMEs perform complex supply chain operations—procurement, logistics, compliance—exactly like SABIC. Yet they lack the data infrastructure, visibility, and intelligence available to large enterprises.

When digitalization was imposed on SMEs, it didn't empower them. It added work, cost, and complexity. Enterprise systems (SAP, Oracle) were built for enterprises, not SMEs. AI platforms built APIs for structured data sources, bypassing SMEs entirely.

Result: 600K+ Saudi SMEs waste 15-30% of time on manual document processing, lack decision visibility, and cannot scale efficiently.

Kraftd solves this.''')

# The Solution
heading = doc_summary.add_heading('What Kraftd Does', level=2)
for run in heading.runs:
    run.font.color.rgb = RGBColor(25, 55, 109)
    run.font.name = 'Segoe UI'

doc_summary.add_paragraph('Convert unstructured documents (invoices, quotations, contracts, BOQs) to structured intelligence (95%+ accuracy)', style='List Bullet')
doc_summary.add_paragraph('Monitor AI tool reliability in organizational workflows (unique capability)', style='List Bullet')
doc_summary.add_paragraph('Enable SMEs to operate with enterprise-grade clarity—without enterprise-grade cost ($500-5K/month vs $50K+)', style='List Bullet')
doc_summary.add_paragraph('2-week implementation vs 6-18 months for traditional ERP', style='List Bullet')

# Traction
heading = doc_summary.add_heading('Traction', level=2)
for run in heading.runs:
    run.font.color.rgb = RGBColor(25, 55, 109)
    run.font.name = 'Segoe UI'

doc_summary.add_paragraph('MVP live January 2026 | 47+ orgs interested | 20+ trials | 95%+ accuracy proven', style='List Bullet')
doc_summary.add_paragraph('Unit economics: LTV:CAC = 13:1 to 25:1 | Payback = 2-4 months | Year 1 path to profitability', style='List Bullet')
doc_summary.add_paragraph('Founder: 12+ years supply chain + SABIC/NEOM experience + built entire MVP solo', style='List Bullet')

# Financials
heading = doc_summary.add_heading('Financial Vision', level=2)
for run in heading.runs:
    run.font.color.rgb = RGBColor(25, 55, 109)
    run.font.name = 'Segoe UI'

doc_summary.add_paragraph('Y1: 15-25 customers, $180-300K revenue | Y2: 50-80 customers, $600K-1.2M, +EBITDA | Y5: 500-1K customers, $6-12M revenue, 40% EBITDA, $50M+ valuation', style='List Bullet')

# Why MISA
heading = doc_summary.add_heading('Strategic Value for MISA', level=2)
for run in heading.runs:
    run.font.color.rgb = RGBColor(25, 55, 109)
    run.font.name = 'Segoe UI'

doc_summary.add_paragraph('Success story proving Saudi Arabia builds world-class SaaS', style='List Bullet')
doc_summary.add_paragraph('Direct impact on 600K+ SMEs + Vision 2030 digitalization', style='List Bullet')
doc_summary.add_paragraph('25-50x ROI probability within 5-7 years', style='List Bullet')

# Closing
closing = doc_summary.add_paragraph('''Kraftd isn't just software—it's rebalancing the intelligence equation in supply chains. With MISA's support, this becomes the proof point for Saudi Arabia's AI/SaaS capability globally.''')
for run in closing.runs:
    run.font.italic = True
    run.font.size = Pt(10)

doc_summary.add_paragraph()
doc_summary.add_paragraph('Contact: akram@kraftd.io | Demo: https://kraftd-a4gfhqa2axb2h6cd.uaenorth-01.azurewebsites.net')

doc_summary.save(r'c:\Users\1R6\OneDrive\Desktop\Kraftd Docs\KRAFTD_One_Page_Summary_MISA.docx')
print("✓ One-Page Summary (Word) created")

print("\n" + "="*60)
print("SUCCESS: Word documents created!")
print("="*60)
print("1. KRAFTD_Company_Profile_MISA.docx")
print("2. KRAFTD_One_Page_Summary_MISA.docx")
