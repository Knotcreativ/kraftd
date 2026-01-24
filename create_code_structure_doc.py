from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def add_heading_styled(doc, text, level, color=(25, 55, 109)):
    """Add styled heading"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(*color)
        run.font.name = 'Segoe UI'
    heading.space_before = Pt(12)
    heading.space_after = Pt(6)
    return heading

def set_cell_background(cell, fill_color):
    """Set cell background color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill_color)
    cell._element.get_or_add_tcPr().append(shading_elm)

def add_code_block(doc, code_text, language="python"):
    """Add a code block with syntax highlighting style"""
    p = doc.add_paragraph()
    p.style = 'Normal'
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0, 100, 0)  # Dark green for code
    return p

# ==================== WORD - CODE STRUCTURE ANALYSIS ====================
doc = Document()
doc.default_tab_stop = Inches(0.5)

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Segoe UI'
font.size = Pt(11)

# Title
title = doc.add_paragraph()
title_run = title.add_run('KRAFTD DOCS')
title_run.font.size = Pt(32)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(25, 55, 109)
title_run.font.name = 'Segoe UI'
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
title.space_after = Pt(2)

subtitle = doc.add_paragraph()
subtitle_run = subtitle.add_run('Complete Code Structure Analysis')
subtitle_run.font.size = Pt(18)
subtitle_run.font.italic = True
subtitle_run.font.color.rgb = RGBColor(100, 100, 100)
subtitle_run.font.name = 'Segoe UI'
subtitle.space_after = Pt(8)

date_line = doc.add_paragraph()
date_run = date_line.add_run('January 23, 2026')
date_run.font.size = Pt(10)
date_run.font.color.rgb = RGBColor(150, 150, 150)
date_run.font.name = 'Segoe UI'

# Executive Summary
add_heading_styled(doc, 'Executive Summary', 1)

summary_text = doc.add_paragraph()
summary_run = summary_text.add_run(
    'This document provides a comprehensive analysis of the Kraftd Docs SaaS platform codebase, '
    'a production-ready AI-powered document intelligence solution serving 600K+ Saudi SMEs in a $400B procurement market. '
    'The platform demonstrates enterprise-grade architecture with sophisticated AI integration, robust security, '
    'and scalable cloud infrastructure.'
)
summary_run.font.name = 'Segoe UI'
summary_run.font.size = Pt(11)

# Architecture Overview
add_heading_styled(doc, 'Architecture Overview', 1)

arch_text = doc.add_paragraph()
arch_run = arch_text.add_run(
    'Kraftd Docs is a production-ready AI-powered document intelligence SaaS platform that transforms '
    'procurement documents (invoices, purchase orders, contracts) into structured data using Azure AI services. '
    'The platform serves 600K+ Saudi SMEs in a $400B procurement market.'
)
arch_run.font.name = 'Segoe UI'

add_heading_styled(doc, 'Core Architecture', 2)
core_arch = doc.add_paragraph()
core_run = core_arch.add_run(
    '• Backend: FastAPI (Python 3.13.9) with async architecture\n'
    '• Frontend: React 18 + TypeScript + Vite (production deployment)\n'
    '• Database: Azure Cosmos DB with multi-tenant partitioning\n'
    '• AI Services: Azure OpenAI GPT-4 + Azure Document Intelligence (95%+ accuracy)\n'
    '• Infrastructure: Azure Container Apps + Azure Static Web Apps\n'
    '• Deployment: Docker containerization + GitHub Actions CI/CD'
)
core_run.font.name = 'Segoe UI'

# Directory Structure Analysis
add_heading_styled(doc, 'Directory Structure Analysis', 1)

add_heading_styled(doc, 'Root Level (/)', 2)
root_text = doc.add_paragraph()
root_run = root_text.add_run(
    '• 18 Route Modules in backend/routes/ - Modular API endpoints\n'
    '• 30+ Service Classes in backend/services/ - Business logic layer\n'
    '• 20+ React Components in frontend/src/components/ - UI building blocks\n'
    '• 15+ Application Pages in frontend/src/pages/ - Route-based views\n'
    '• Infrastructure Templates in infrastructure/ - Azure Bicep deployments\n'
    '• 50+ Documentation Files in docs/ - Comprehensive system docs\n'
    '• 40+ Automation Scripts in scripts/ - Deployment & testing tools'
)
root_run.font.name = 'Segoe UI'

add_heading_styled(doc, 'Backend Architecture (backend/)', 2)
backend_arch = doc.add_paragraph()
backend_run = backend_arch.add_run(
    'backend/\n'
    '├── main.py (2,425 lines) - FastAPI app with 18 route modules\n'
    '├── routes/ (16 modules) - API endpoint definitions\n'
    '├── services/ (30+ classes) - Business logic & integrations\n'
    '├── models/ - Pydantic data models & error handling\n'
    '├── repositories/ - Cosmos DB data access layer\n'
    '├── middleware/ - Authentication, RBAC, rate limiting\n'
    '├── document_processing/ - AI document extraction pipeline\n'
    '├── agent/ - GPT-4o mini AI agent integration\n'
    '├── ml/ - Machine learning models & predictions\n'
    '├── utils/ - Helper functions & utilities\n'
    '├── tests/ - 332+ unit tests (pytest)\n'
    '└── requirements.txt (40+ dependencies)'
)
backend_run.font.name = 'Consolas'
backend_run.font.size = Pt(9)

add_heading_styled(doc, 'Frontend Architecture (frontend/)', 2)
frontend_arch = doc.add_paragraph()
frontend_run = frontend_arch.add_run(
    'frontend/\n'
    '├── src/\n'
    '│   ├── components/ (20+ React components)\n'
    '│   │   ├── DocumentUpload.tsx - File upload interface\n'
    '│   │   ├── Dashboard.tsx - Analytics dashboard\n'
    '│   │   ├── AnalyticsCharts.tsx - Data visualization\n'
    '│   │   └── Auth components - Login/register flows\n'
    '│   ├── pages/ (15+ route pages)\n'
    '│   │   ├── Login.tsx, Register.tsx - Authentication\n'
    '│   │   ├── Dashboard.tsx - Main application view\n'
    '│   │   ├── DocumentDetail.tsx - Document viewer\n'
    '│   │   └── Admin panels - User management\n'
    '│   ├── services/ - API client & utilities\n'
    '│   ├── context/ - React context providers\n'
    '│   ├── hooks/ - Custom React hooks\n'
    '│   └── types/ - TypeScript interfaces\n'
    '├── package.json (15 dependencies)\n'
    '└── staticwebapp.config.json - Azure deployment config'
)
frontend_run.font.name = 'Consolas'
frontend_run.font.size = Pt(9)

# Core Technologies & Dependencies
add_heading_styled(doc, 'Core Technologies & Dependencies', 1)

add_heading_styled(doc, 'Backend Dependencies (40+ packages)', 2)
backend_deps = doc.add_paragraph()
backend_deps_run = backend_deps.add_run(
    '• Web Framework: fastapi, uvicorn\n'
    '• AI/ML: openai, azure-ai-documentintelligence\n'
    '• Database: azure-cosmos, psycopg2-binary\n'
    '• Security: PyJWT, passlib[bcrypt], azure-identity\n'
    '• Data Processing: pandas, scikit-learn, numpy\n'
    '• Document Processing: pdfplumber, python-docx, openpyxl, pytesseract, pillow\n'
    '• Async: httpx, aiofiles\n'
    '• Testing: pytest, pytest-asyncio, pytest-cov\n'
    '• Utilities: python-dotenv, email-validator, sendgrid'
)
backend_deps_run.font.name = 'Segoe UI'

add_heading_styled(doc, 'Frontend Dependencies (15 packages)', 2)
frontend_deps = doc.add_paragraph()
frontend_deps_run = frontend_deps.add_run(
    '• Core: react, react-dom, react-router-dom\n'
    '• Build: @vitejs/plugin-react, typescript, vite\n'
    '• UI: react-beautiful-dnd, react-datepicker, recharts\n'
    '• HTTP: axios\n'
    '• File Handling: file-saver, xlsx\n'
    '• Dev Tools: TypeScript types for all dependencies'
)
frontend_deps_run.font.name = 'Segoe UI'

# API Architecture
add_heading_styled(doc, 'API Architecture (18 Route Modules)', 1)

add_heading_styled(doc, 'Core API Routes (/api/v1/)', 2)
api_routes = doc.add_paragraph()
api_routes_run = api_routes.add_run(
    '1. /auth - User authentication & authorization\n'
    '2. /agent - AI agent chat & document intelligence\n'
    '3. /documents - Document upload, management & retrieval\n'
    '4. /conversions - Document conversion sessions\n'
    '5. /extraction - Data extraction & processing\n'
    '6. /exports - Output generation & feedback\n'
    '7. /schema - Schema generation & validation\n'
    '8. /user - User profile management\n'
    '9. /ml/advanced - Advanced analytics & ML predictions\n'
    '10. /templates - Document template generation\n'
    '11. /signals - Real-time signals intelligence\n'
    '12. /events - Historical data & analytics\n'
    '13. /data-enhancement - AI-enhanced data processing\n'
    '14. /streaming - WebSocket real-time updates\n'
    '15. /quota - Usage limits & billing\n'
    '16. /admin - Administrative functions'
)
api_routes_run.font.name = 'Segoe UI'

add_heading_styled(doc, 'Route Registration Pattern', 2)
route_code = doc.add_paragraph()
route_code_run = route_code.add_run(
    '# Conditional route inclusion with availability checks\n'
    'if AUTH_ROUTES_AVAILABLE:\n'
    '    app.include_router(auth_router, prefix="/api/v1")\n'
    'if DOCUMENTS_ROUTES_AVAILABLE:\n'
    '    app.include_router(documents_router, prefix="/api/v1")\n'
    '# ... 14 more route modules'
)
route_code_run.font.name = 'Consolas'
route_code_run.font.size = Pt(9)

# Data Architecture
add_heading_styled(doc, 'Data Architecture', 1)

add_heading_styled(doc, 'Cosmos DB Multi-Tenant Design', 2)
cosmos_design = doc.add_paragraph()
cosmos_design_run = cosmos_design.add_run(
    '• Partition Key: user_email for tenant isolation\n'
    '• Containers: documents, users, conversions, extractions\n'
    '• TTL: 30 days for temporary data, indefinite for user data\n'
    '• Indexing: Consistent indexing with included/excluded paths\n'
    '• Consistency: Session-level consistency for performance'
)
cosmos_design_run.font.name = 'Segoe UI'

add_heading_styled(doc, 'Repository Pattern', 2)
repo_pattern = doc.add_paragraph()
repo_pattern_run = repo_pattern.add_run(
    '# Async repository classes for data access\n'
    'class DocumentRepository:\n'
    '    async def create_document(self, doc: dict, user_email: str)\n'
    '    async def get_document(self, doc_id: str, user_email: str)\n'
    '    async def update_document(self, doc_id: str, user_email: str, updates: dict)'
)
repo_pattern_run.font.name = 'Consolas'
repo_pattern_run.font.size = Pt(9)

# AI Integration Architecture
add_heading_styled(doc, 'AI Integration Architecture', 1)

add_heading_styled(doc, 'Document Intelligence Pipeline', 2)
ai_pipeline = doc.add_paragraph()
ai_pipeline_run = ai_pipeline.add_run(
    '1. Azure Document Intelligence - 95%+ accuracy OCR/extraction\n'
    '2. GPT-4o Mini Agent - Intelligent document analysis\n'
    '3. ML Models - Predictive analytics & pattern recognition\n'
    '4. Quality Validation - Confidence scoring & error detection'
)
ai_pipeline_run.font.name = 'Segoe UI'

add_heading_styled(doc, 'Agent Framework', 2)
agent_framework = doc.add_paragraph()
agent_framework_run = agent_framework.add_run(
    'class KraftdAIAgent:\n'
    '    async def process_message(self, message: str, conversation_id: str)\n'
    '    async def get_learning_insights(self)\n'
    '    async def _sync_learning_patterns(self)'
)
agent_framework_run.font.name = 'Consolas'
agent_framework_run.font.size = Pt(9)

# Security & Authentication
add_heading_styled(doc, 'Security & Authentication', 1)

add_heading_styled(doc, 'Multi-Layer Security', 2)
security_layers = doc.add_paragraph()
security_layers_run = security_layers.add_run(
    '• JWT Authentication with refresh token rotation\n'
    '• RBAC (Role-Based Access Control) with granular permissions\n'
    '• Rate Limiting (requests/minute & requests/hour)\n'
    '• CORS Configuration with origin restrictions\n'
    '• Audit Logging for compliance tracking\n'
    '• Multi-Tenant Isolation via partition keys'
)
security_layers_run.font.name = 'Segoe UI'

add_heading_styled(doc, 'Middleware Stack', 2)
middleware_stack = doc.add_paragraph()
middleware_stack_run = middleware_stack.add_run(
    'app.add_middleware(CORSMiddleware, allow_origins=cors_origins)\n'
    'app.add_middleware(RateLimitMiddleware, requests_per_minute=100)\n'
    '# RBAC middleware applied to protected routes'
)
middleware_stack_run.font.name = 'Consolas'
middleware_stack_run.font.size = Pt(9)

# Monitoring & Observability
add_heading_styled(doc, 'Monitoring & Observability', 1)

add_heading_styled(doc, 'Built-in Monitoring', 2)
monitoring_features = doc.add_paragraph()
monitoring_features_run = monitoring_features.add_run(
    '• Application Insights integration\n'
    '• Custom Metrics Collection (metrics.py)\n'
    '• Health Checks (/api/v1/health)\n'
    '• Error Tracking with structured logging\n'
    '• Performance Monitoring with request timing\n'
    '• Export Tracking (three-stage recording)'
)
monitoring_features_run.font.name = 'Segoe UI'

# Deployment Architecture
add_heading_styled(doc, 'Deployment Architecture', 1)

add_heading_styled(doc, 'Production Infrastructure', 2)
prod_infra = doc.add_paragraph()
prod_infra_run = prod_infra.add_run(
    '• Backend: Azure Container Apps (Docker containerized)\n'
    '• Frontend: Azure Static Web Apps (Vite-built SPA)\n'
    '• Database: Azure Cosmos DB (serverless)\n'
    '• AI Services: Azure Cognitive Services\n'
    '• CDN: Azure CDN for global distribution'
)
prod_infra_run.font.name = 'Segoe UI'

add_heading_styled(doc, 'CI/CD Pipeline', 2)
cicd_pipeline = doc.add_paragraph()
cicd_pipeline_run = cicd_pipeline.add_run(
    '• GitHub Actions for automated deployment\n'
    '• Docker Build with multi-stage optimization\n'
    '• Azure CLI Integration for infrastructure provisioning\n'
    '• Environment Configuration via .env files'
)
cicd_pipeline_run.font.name = 'Segoe UI'

# Testing & Quality Assurance
add_heading_styled(doc, 'Testing & Quality Assurance', 1)

add_heading_styled(doc, 'Test Coverage', 2)
test_coverage = doc.add_paragraph()
test_coverage_run = test_coverage.add_run(
    '• 332 Unit Tests passing with multi-tenant validation\n'
    '• pytest Framework with async support\n'
    '• Integration Tests for API endpoints\n'
    '• E2E Testing scripts for full workflow validation\n'
    '• Load Testing capabilities'
)
test_coverage_run.font.name = 'Segoe UI'

# Development Workflow
add_heading_styled(doc, 'Development Workflow', 1)

add_heading_styled(doc, 'Local Development', 2)
local_dev = doc.add_paragraph()
local_dev_run = local_dev.add_run(
    '# Backend\n'
    'cd backend && uvicorn main:app --reload\n'
    '\n'
    '# Frontend\n'
    'cd frontend && npm run dev\n'
    '\n'
    '# Testing\n'
    'cd backend && python -m pytest'
)
local_dev_run.font.name = 'Consolas'
local_dev_run.font.size = Pt(9)

add_heading_styled(doc, 'Production Deployment', 2)
prod_deploy = doc.add_paragraph()
prod_deploy_run = prod_deploy.add_run(
    '# Automated deployment\n'
    './deploy-azure.bat\n'
    '\n'
    '# Manual steps\n'
    'az login\n'
    'az account set --subscription <subscription-id>\n'
    'az deployment group create --resource-group kraftd-docs-rg --template-file infrastructure/main.bicep'
)
prod_deploy_run.font.name = 'Consolas'
prod_deploy_run.font.size = Pt(9)

# Scaling & Performance
add_heading_styled(doc, 'Scaling & Performance', 1)

add_heading_styled(doc, 'Performance Optimizations', 2)
perf_opts = doc.add_paragraph()
perf_opts_run = perf_opts.add_run(
    '• Async/Await throughout the application\n'
    '• Connection Pooling for database connections\n'
    '• Caching for frequently accessed data\n'
    '• Rate Limiting to prevent abuse\n'
    '• Horizontal Scaling via Azure Container Apps'
)
perf_opts_run.font.name = 'Segoe UI'

# Business Impact
add_heading_styled(doc, 'Business Impact', 1)

add_heading_styled(doc, 'Market Position', 2)
market_pos = doc.add_paragraph()
market_pos_run = market_pos.add_run(
    '• Target Market: 600K+ Saudi SMEs, $400B procurement market\n'
    '• Competitive Advantage: 100x cheaper than enterprise ERP, 2-week implementation\n'
    '• Financial Projections: Year 1 $180-300K, Year 5 $6-12M, LTV:CAC 13:1-25:1'
)
market_pos_run.font.name = 'Segoe UI'

add_heading_styled(doc, 'Technical Differentiation', 2)
tech_diff = doc.add_paragraph()
tech_diff_run = tech_diff.add_run(
    '• AI Accuracy: 95%+ document extraction accuracy\n'
    '• Multi-Format Support: PDF, Word, Excel, images\n'
    '• Real-Time Processing: WebSocket streaming for live updates\n'
    '• Multi-Tenant Security: Complete data isolation per user'
)
tech_diff_run.font.name = 'Segoe UI'

# Conclusion
add_heading_styled(doc, 'Conclusion', 1)

conclusion_text = doc.add_paragraph()
conclusion_run = conclusion_text.add_run(
    'This comprehensive analysis reveals a production-ready, enterprise-grade SaaS platform with sophisticated AI integration, '
    'robust security, and scalable cloud architecture. The codebase demonstrates advanced software engineering practices with '
    '332+ tests, comprehensive documentation, and automated deployment pipelines. The platform is fully prepared for market '
    'launch and enterprise scaling.'
)
conclusion_run.font.name = 'Segoe UI'

# Save the document
output_path = r'c:\Users\1R6\OneDrive\Desktop\Kraftd Docs\KRAFTD_Code_Structure_Analysis.docx'
doc.save(output_path)

print(f"Code Structure Analysis document created successfully at: {output_path}")
print("Opening document...")

# Open the document
os.startfile(output_path)