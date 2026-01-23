from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Define MISA colors
PRIMARY = RGBColor(25, 55, 109)      # Deep blue
ACCENT = RGBColor(0, 150, 136)       # Teal
TEXT = RGBColor(50, 50, 50)          # Dark gray
WHITE = RGBColor(255, 255, 255)      # White

# Load existing document
doc = Document(r'c:\Users\1R6\OneDrive\Desktop\Kraftd Docs\Kraftd Company Profile copy.docx')

# Apply formatting to all paragraphs
for para in doc.paragraphs:
    # Set font to Segoe UI for all text in paragraph
    for run in para.runs:
        run.font.name = 'Segoe UI'
        run.font.size = Pt(11)
        
        # Apply color based on paragraph style/content
        if para.style.name.startswith('Heading'):
            # Headings get primary color and bold
            run.font.color.rgb = PRIMARY
            run.font.bold = True
            run.font.size = Pt(16) if 'Heading 1' in para.style.name else Pt(13)
        else:
            # Body text gets dark gray
            run.font.color.rgb = TEXT
    
    # Set paragraph alignment
    if para.style.name.startswith('Heading'):
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT

# Format tables if they exist
for table in doc.tables:
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            # Header row - primary color background with white text
            if i == 0:
                # Set cell fill color
                from docx.oxml import parse_xml
                from docx.oxml.ns import nsdecls
                
                shading_elm = parse_xml(r'<w:shd {} w:fill="193D6D"/>'.format(nsdecls('w')))
                cell._element.get_or_add_tcPr().append(shading_elm)
                
                # Format text in header
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Segoe UI'
                        run.font.bold = True
                        run.font.size = Pt(11)
                        run.font.color.rgb = WHITE
            else:
                # Data rows - light background
                shading_elm = parse_xml(r'<w:shd {} w:fill="F0F7FF"/>'.format(nsdecls('w')))
                cell._element.get_or_add_tcPr().append(shading_elm)
                
                # Format text in data cells
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Segoe UI'
                        run.font.size = Pt(10)
                        run.font.color.rgb = TEXT

# Save with MISA template applied
output_path = r'c:\Users\1R6\OneDrive\Desktop\Kraftd Docs\Kraftd_Company_Profile_MISA_Template.docx'
doc.save(output_path)

print("✓ Template applied successfully!")
print(f"\nFormatting Applied:")
print("=" * 60)
print(f"Font: Segoe UI (all content)")
print(f"Headings: Primary Blue (RGB 25, 55, 109), Bold, 16pt")
print(f"Body Text: Dark Gray (RGB 50, 50, 50), 11pt")
print(f"Table Headers: Primary Blue background, White text")
print(f"Table Data: Light blue background, Dark text")
print("=" * 60)
print(f"\nOutput saved to:")
print(f"{output_path}")
print("\nContent preserved - only formatting applied")
