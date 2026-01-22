#!/usr/bin/env python3
"""
Generate KRAFTD Business Analysis Word Document
Converts manifesto-centered analysis to professional .docx format
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_heading(doc, text, level=1):
    """Add formatted heading"""
    heading = doc.add_heading(text, level=level)
    heading.paragraph_format.space_before = Pt(12)
    heading.paragraph_format.space_after = Pt(6)
    return heading

def add_paragraph(doc, text, bold=False, italic=False, color=None):
    """Add formatted paragraph"""
    para = doc.add_paragraph(text)
    if bold or italic or color:
        for run in para.runs:
            if bold:
                run.font.bold = True
            if italic:
                run.font.italic = True
            if color:
                run.font.color.rgb = color
    return para

def set_cell_background(cell, fill):
    """Set cell background color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill)
    cell._element.get_or_add_tcPr().append(shading_elm)

# Create document
doc = Document()
doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(11)

# Title Page
title = doc.add_heading('KRAFTD', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_paragraph('Intelligence Equity in Supply Chains')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(18)
subtitle.runs[0].font.bold = True

tagline = doc.add_paragraph('Investment Thesis Centered on Structural Change')
tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
tagline.runs[0].font.italic = True
tagline.runs[0].font.size = Pt(12)

doc.add_paragraph()
doc.add_paragraph()

info = doc.add_paragraph('January 20, 2026')
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.runs[0].font.size = Pt(11)

doc.add_page_break()

# TABLE OF CONTENTS
add_heading(doc, 'TABLE OF CONTENTS', 1)
toc_items = [
    'MANIFESTO',
    'THE OPPORTUNITY: Supply Chain Intelligence Equity',
    'WHY EXISTING SOLUTIONS FAIL',
    'THE KRAFTD SOLUTION: Translation Layer',
    'MARKET POSITIONING',
    'THE BUSINESS MODEL',
    'FINANCIAL PROJECTIONS',
    'COMPETITIVE ANALYSIS',
    'SWOT ANALYSIS (Revised)',
    'FEASIBILITY STUDY',
    'INVESTMENT OPPORTUNITY',
    'EXECUTIVE SUMMARY FOR INVESTORS',
]

for item in toc_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# MANIFESTO SECTION
add_heading(doc, 'MANIFESTO', 1)

add_heading(doc, 'The Structural Problem', 2)
doc.add_paragraph('Modern supply chains run on intelligence.')
doc.add_paragraph('But that intelligence is unevenly distributed.')
doc.add_paragraph()

doc.add_paragraph('Large enterprises operate with structured systems, integrated ERPs, APIs, dashboards, and predictive analytics. Their data is clean, connected, and continuously exploited for advantage.')
doc.add_paragraph()

doc.add_paragraph('SMEs and suppliers — who extract, process, transport, and deliver the same goods — operate in the dark.')
doc.add_paragraph()

doc.add_paragraph('This is not a talent gap.')
para = doc.add_paragraph('It is a systems design failure.')
para.runs[0].font.italic = True

add_heading(doc, 'The Architecture of Inequality', 2)
doc.add_paragraph('The global AI ecosystem is built around APIs.')
doc.add_paragraph()
doc.add_paragraph('APIs are cheap, scalable, and elegant—but only if you already have structured data.')
doc.add_paragraph()
doc.add_paragraph('That assumption quietly excludes most of the world\'s productive economy.')
doc.add_paragraph()
doc.add_paragraph('SMEs do not live in schemas.')
doc.add_paragraph('They live in PDFs, invoices, contracts, emails, spreadsheets, drawings, customs forms, portals, and manual reconciliation.')
doc.add_paragraph()

doc.add_paragraph('Every day, SMEs spend hours reshaping raw reality into formats demanded by enterprise systems:')
for item in ['Normalizing data', 'Re-entering information', 'Reconciling inconsistencies', 'Absorbing compliance risk', 'Subsidizing the intelligence of others']:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph()
doc.add_paragraph('Meanwhile, they receive no equivalent visibility, foresight, or decision support in return.')
doc.add_paragraph()
doc.add_paragraph('The system was never built for them.')

add_heading(doc, 'What This Means', 2)
doc.add_paragraph('An SME performing identical supply chain operations as SABIC has:')
items_negative = ['No structured data visibility', 'No API access to market intelligence', 'No predictive analytics for risk', 'No supplier visibility tools', 'No decision support equivalent to enterprise platforms']
for item in items_negative:
    doc.add_paragraph('✗ ' + item, style='List Bullet')

doc.add_paragraph()
doc.add_paragraph('Yet they:')
items_positive = ['Operate the same supply chain steps', 'Face the same risks', 'Must meet the same compliance burdens', 'Generate the same data (just messier)']
for item in items_positive:
    doc.add_paragraph('✓ ' + item, style='List Bullet')

doc.add_paragraph()
para = doc.add_paragraph('This is not inevitable. This is a design choice.')
para.runs[0].font.bold = True

add_heading(doc, 'Our Belief', 2)
doc.add_paragraph('Intelligence should not require ERP maturity.')
doc.add_paragraph('Insight should not require API readiness.')
doc.add_paragraph('Credibility should not require scale.')
doc.add_paragraph()
doc.add_paragraph('If an SME can operate a business, it should be able to operate with intelligence.')

doc.add_page_break()

# OPPORTUNITY SECTION
add_heading(doc, 'THE OPPORTUNITY: Supply Chain Intelligence Equity', 1)

add_heading(doc, 'Market Reality', 2)
doc.add_paragraph('Global supply chains = $118 trillion annually', style='List Bullet')
doc.add_paragraph('Large enterprises: Well-served, structured data, AI/ML platforms everywhere', style='List Bullet')
doc.add_paragraph('SME & supplier ecosystem: $35-50 trillion, <15% digitalized, zero intelligence parity', style='List Bullet')

doc.add_paragraph()
para = doc.add_paragraph('This is not "the SME procurement market."')
para.runs[0].font.bold = True
doc.add_paragraph()
doc.add_paragraph('This is supply chain segments where intelligence has never reached:')
doc.add_paragraph()

segments = [
    ('Agriculture Trading Platforms', ['Farmer suppliers → No visibility', 'Processing partners → Submitting to portals', 'Enterprise buyers → Full supply chain visibility']),
    ('Manufacturing Supply Chains', ['Raw material suppliers → Manual tracking', 'Sub-component makers → Spreadsheet reconciliation', 'Contract manufacturers → Compliance-heavy portals', 'Brand enterprises → Full visibility dashboard']),
    ('Logistics Networks', ['Regional 3PLs → Manual capacity tracking', 'Local carriers → Portal-based compliance', 'Customs brokers → Document shuffling', 'Global corporations → Real-time visibility']),
    ('Construction & Project Supply', ['Material suppliers → Invoice tracking', 'Sub-contractors → Portal uploads', 'Equipment providers → Compliance forms', 'Developers → Full supply transparency']),
    ('Commodity Trading', ['Producers → Manual quotations', 'Traders → Spreadsheet analysis', 'Exporters → Document management', 'Importers/Buyers → Full market data access']),
]

for segment, details in segments:
    para = doc.add_paragraph(segment)
    para.runs[0].font.bold = True
    for detail in details:
        doc.add_paragraph(detail, style='List Bullet')
    doc.add_paragraph()

add_heading(doc, 'Addressable Market: NOT Just Procurement', 2)
doc.add_paragraph('Initial analysis focused on "procurement efficiency for SMEs" = ~$550B addressable market')
doc.add_paragraph()
para = doc.add_paragraph('The manifesto reframes it:')
para.runs[0].font.bold = True
doc.add_paragraph()
para = doc.add_paragraph('The real market is "Supply chain intelligence for fragmented operators" = $2-3 trillion+ globally')
para.runs[0].font.bold = True

# Add market sizing table
table = doc.add_table(rows=7, cols=2)
table.style = 'Light Grid Accent 1'

# Header row
header_cells = table.rows[0].cells
header_cells[0].text = 'Segment'
header_cells[1].text = 'Market Size'

# Data rows
data = [
    ('Procurement', '$550B'),
    ('Supplier Intelligence', '$300B'),
    ('Trading/Import-Export', '$250B'),
    ('Manufacturing/Processing', '$200B'),
    ('Logistics', '$150B'),
    ('TOTAL GCC Region', '$1.45 TRILLION'),
]

for i, (segment, size) in enumerate(data, 1):
    row_cells = table.rows[i].cells
    row_cells[0].text = segment
    row_cells[1].text = size
    if i == 6:  # Last row
        set_cell_background(row_cells[0], 'D3D3D3')
        set_cell_background(row_cells[1], 'D3D3D3')

doc.add_paragraph()
doc.add_paragraph('Global: Fragmented Supply Chain Operators = $2-3 TRILLION')

doc.add_page_break()

# WHY EXISTING SOLUTIONS FAIL
add_heading(doc, 'WHY EXISTING SOLUTIONS FAIL', 1)

add_heading(doc, 'SAP, Oracle, Coupa Problem', 2)
doc.add_paragraph('These platforms are architecturally incapable of serving messy-data SMEs because:')
doc.add_paragraph()

problems = [
    ('Built for Clean Schemas', ['Assume data arrives structured', 'Require predefined templates', 'Break with deviation']),
    ('ERP-First Model', ['Assume customer has enterprise infrastructure', 'Prerequisite: $500K-5M implementation', 'Prerequisite: Existing IT team']),
    ('Upward Flow Only', ['Designed to extract supplier data for enterprise use', 'Returns compliance requirements, not intelligence', 'Never built to serve the supplier']),
    ('Economics Don\'t Work', ['Sales/implementation cost = $150K-$1M', 'Only viable for $50M+ companies', 'Margin doesn\'t work for $5M companies']),
]

for problem_title, details in problems:
    para = doc.add_paragraph(problem_title)
    para.runs[0].font.bold = True
    para.runs[0].font.color.rgb = RGBColor(192, 0, 0)
    for detail in details:
        doc.add_paragraph(detail, style='List Bullet')

doc.add_paragraph()
para = doc.add_paragraph('Result: These vendors have structurally abandoned the SME/supplier segment.')
para.runs[0].font.italic = True
doc.add_paragraph()
para = doc.add_paragraph('This isn\'t competition. This is market indifference at scale.')
para.runs[0].font.bold = True

doc.add_page_break()

# THE KRAFTD SOLUTION
add_heading(doc, 'THE KRAFTD SOLUTION: Translation Layer', 1)

doc.add_paragraph('Kraftd is not "better procurement software."')
doc.add_paragraph()
para = doc.add_paragraph('Kraftd is a translation layer between messy reality and intelligence.')
para.runs[0].font.bold = True

doc.add_paragraph()
doc.add_paragraph('We operate in the space enterprise systems cannot reach:')
doc.add_paragraph()

doc.add_paragraph('Before APIs', style='List Bullet')
doc.add_paragraph('↓', style='List Bullet')
doc.add_paragraph('Before Schemas', style='List Bullet')
doc.add_paragraph('↓', style='List Bullet')
doc.add_paragraph('Before Enterprise Infrastructure', style='List Bullet')
doc.add_paragraph('↓', style='List Bullet')
doc.add_paragraph('WHERE DATA LIVES IN REALITY', style='List Bullet')
doc.add_paragraph('↓', style='List Bullet')
para = doc.add_paragraph('[KRAFTD OPERATES HERE]')
para.runs[0].font.bold = True
doc.add_paragraph('↓', style='List Bullet')
doc.add_paragraph('Translates to Actionable Intelligence', style='List Bullet')
doc.add_paragraph('↓', style='List Bullet')
doc.add_paragraph('Supply to Enterprise Systems (if needed)', style='List Bullet')

add_heading(doc, 'What Kraftd Does', 2)

capabilities = [
    ('Translates Messy Documents to Intelligence', 
     'SMEs work with documents. Enterprise systems require structured data.\n\nKraftd bridges this gap:\n• Takes RFQs, invoices, contracts, customs forms, delivery notes (messy)\n• Extracts structured intelligence (95%+ accuracy)\n• Produces insights without requiring data entry'),
    
    ('Makes Intelligence Work Before ERP',
     'Enterprise systems assume you already have infrastructure.\n\nKraftd works without prerequisites:\n• No ERP required\n• No IT team required\n• No pre-existing data structure required\n• No API integration required (though it works with them)'),
    
    ('Converts Compliance Burden into Insight',
     'When enterprises demand supplier data, it creates burden on SMEs.\n\nKraftd inverts this:\n• Supplier uploads documents to meet compliance\n• Kraftd extracts intelligence for the supplier\n• Same effort, but now both parties benefit'),
    
    ('Provides Intelligence Parity',
     'Large enterprises see supply chain risks, opportunities, and patterns.\n\nKraftd gives SMEs the same visibility:\n• Price trends vs market baseline\n• Supplier reliability patterns\n• Risk indicators before problems emerge\n• Competitive positioning intelligence'),
]

for title, description in capabilities:
    para = doc.add_paragraph(title)
    para.runs[0].font.bold = True
    para.runs[0].font.color.rgb = RGBColor(0, 102, 204)
    doc.add_paragraph(description)
    doc.add_paragraph()

doc.add_page_break()

# MARKET POSITIONING
add_heading(doc, 'MARKET POSITIONING', 1)

add_heading(doc, 'NOT Competing with Coupa/Ariba', 2)
doc.add_paragraph('We are not a cheaper alternative to enterprise procurement platforms.')
doc.add_paragraph()
para = doc.add_paragraph('We are solving a different problem for a different tier.')
para.runs[0].font.bold = True

# Positioning table
table = doc.add_table(rows=4, cols=4)
table.style = 'Light Grid Accent 1'

headers = ['Layer', 'Problem', 'Solution', 'Market']
for i, header in enumerate(headers):
    table.rows[0].cells[i].text = header
    set_cell_background(table.rows[0].cells[i], 'D3D3D3')

positioning_data = [
    ('Enterprise', 'Maximize value from clean supply chain data', 'Coupa, Ariba, Jaggr', '$500B+ (well-served)'),
    ('Translation', 'Extract intelligence from messy reality', 'KRAFTD', '$2-3T (ignored)'),
    ('Tactical', 'Survive compliance requirements', 'ERP, portals, spreadsheets', 'Endless (painful)'),
]

for i, (layer, problem, solution, market) in enumerate(positioning_data, 1):
    table.rows[i].cells[0].text = layer
    table.rows[i].cells[1].text = problem
    table.rows[i].cells[2].text = solution
    table.rows[i].cells[3].text = market
    if i == 2:  # KRAFTD row
        for cell in table.rows[i].cells:
            set_cell_background(cell, 'E6F2FF')

add_heading(doc, 'What Makes Kraftd Defensible', 2)

defensibility_points = [
    ('First-Mover in Translation', 'No competitor serves "before the API" segment\nThis space is ignored by incumbents\nStructural advantage (they can\'t pivot without rebuilding)'),
    
    ('Architecture Matters', 'Built for messy documents from day 1\nScaled for SME economics ($5-50K/month, not $500K+ implementation)\nDesigned around "compliance effort = insight production"'),
    
    ('Competitive Moat is Structural', 'Coupa could not compete here without destroying their enterprise product\nThey need schemas; we work with chaos\nThey need ERP integration; we work standalone'),
    
    ('Customer Economics', 'SME willingness to pay: $5-50K/month\nOur cost to serve: $100-500/month (90% margin)\nEnterprise platform model: $500K implementation, requires 3-6 months\nOur model: 2-4 week onboarding, plug-and-play'),
]

for title, details in defensibility_points:
    para = doc.add_paragraph(title)
    para.runs[0].font.bold = True
    doc.add_paragraph(details)
    doc.add_paragraph()

doc.add_page_break()

# THE BUSINESS MODEL
add_heading(doc, 'THE BUSINESS MODEL', 1)

add_heading(doc, 'How Intelligence Equity Creates Sustainable Revenue', 2)

add_heading(doc, 'Revenue Stream 1: SME Intelligence (70% of revenue)', 3)

tiers = [
    ('Tier 1: Supplier Intelligence - $499/month', 
     'Document analysis (50 docs/month)\nSupplier risk scoring\nPrice benchmarking\nCompliance tracking\nTarget: 500 customers Year 1 → $3M ARR'),
    
    ('Tier 2: Trade Intelligence - $2,499/month',
     'Multi-party visibility (200 docs/month)\nSupply chain transparency\nCustoms/logistics intelligence\nMarket trend analysis\nTarget: 50 customers Year 1 → $1.5M ARR'),
    
    ('Tier 3: Strategic Intelligence - $9,999/month',
     'Full supply chain orchestration\nPredictive risk modeling\nCompetitive positioning analysis\nSupplier ecosystem mapping\nTarget: 10 customers Year 1 → $1.2M ARR'),
]

for tier_title, tier_details in tiers:
    para = doc.add_paragraph(tier_title)
    para.runs[0].font.bold = True
    doc.add_paragraph(tier_details)

add_heading(doc, 'Revenue Stream 2: Enterprise Supply Chain Transparency (20% of revenue)', 3)

doc.add_paragraph('Enterprises benefit when suppliers have intelligence:')
benefits = ['Better data quality from suppliers (compliance = insight)', 'Faster response from suppliers (they see issues early)', 'Reduced supply chain risk (SME visibility = early warning)', 'Improved supplier relationships (fairness = trust)']
for benefit in benefits:
    doc.add_paragraph(benefit, style='List Bullet')

doc.add_paragraph()
doc.add_paragraph('Model: Enterprise pays for supplier visibility package\n• Kraft provides suppliers with intelligence\n• Enterprise gets transparency dashboard\n• Revenue: $50K-$500K per enterprise\n• Target: 5 customers Year 1 → $500K')

add_heading(doc, 'Revenue Stream 3: Supply Chain Intelligence Data (10% of revenue)', 3)

doc.add_paragraph('As we aggregate SME data, we create new intelligence products:')
data_products = ['Price indices by commodity, region, supplier type', 'Supply chain risk signals (early warning systems)', 'Supplier reliability reports (anonymized, aggregated)', 'Trade opportunity maps (buyer-seller matching)']
for product in data_products:
    doc.add_paragraph(product, style='List Bullet')

doc.add_paragraph()
doc.add_paragraph('Year 3-5 potential: $1M+/year as data moat builds')

doc.add_page_break()

# FINANCIAL PROJECTIONS
add_heading(doc, 'FINANCIAL PROJECTIONS', 1)

add_heading(doc, '5-Year Model Centered on Intelligence Equity', 2)

# Financial summary table
table = doc.add_table(rows=6, cols=8)
table.style = 'Light Grid Accent 1'

headers = ['Metric', 'Year 1', 'Year 2', 'Year 3', 'Year 4', 'Year 5', 'Status', '']
for i, header in enumerate(headers[:6]):
    table.rows[0].cells[i].text = header
    set_cell_background(table.rows[0].cells[i], 'D3D3D3')

financial_data = [
    ('SME Customers', '560', '1,200', '2,500', '3,700', '5,000+'),
    ('Total ARR', '$6M', '$18.5M', '$58.5M', '$90M+', '$120M+'),
    ('Total Revenue', '$6.4M', '$19.3M', '$60M', '$100M+', '$140M+'),
    ('EBITDA', '$560K', '$8M', '$34M', '$55M+', '$80M+'),
    ('Margin', '9%', '41%', '57%', '55%', '57%'),
]

for i, row_data in enumerate(financial_data, 1):
    table.rows[i].cells[0].text = row_data[0]
    for j, value in enumerate(row_data[1:], 1):
        table.rows[i].cells[j].text = value

doc.add_paragraph()
doc.add_paragraph('✓ PROFITABLE BY YEAR 1')

add_heading(doc, 'Unit Economics (Based on Intelligence Model)', 3)

# Unit economics
doc.add_paragraph('CUSTOMER ACQUISITION:')
doc.add_paragraph()

unit_tiers = [
    ('Tier 1 SME (Supplier Intelligence)',
     'CAC: $2,000\nLTV: $9,000 (18 months @ $499, <5% churn)\nLTV:CAC: 4.5x ✓'),
    
    ('Tier 2 SME (Trade Intelligence)',
     'CAC: $5,000\nLTV: $45,000 (18 months @ $2,499, <5% churn)\nLTV:CAC: 9x ✓'),
    
    ('Tier 3 SME (Strategic Intelligence)',
     'CAC: $15,000\nLTV: $180,000 (18 months @ $9,999, <5% churn)\nLTV:CAC: 12x ✓'),
    
    ('Enterprise Transparency',
     'CAC: $50,000\nLTV: $1M+ (3-5 year contracts)\nLTV:CAC: 20x+ ✓'),
]

for tier_name, tier_metrics in unit_tiers:
    para = doc.add_paragraph(tier_name)
    para.runs[0].font.bold = True
    doc.add_paragraph(tier_metrics)

doc.add_paragraph()
doc.add_paragraph('BLENDED METRICS:')
doc.add_paragraph('Blended CAC: $6,000', style='List Bullet')
doc.add_paragraph('Blended LTV: $50,000', style='List Bullet')
doc.add_paragraph('Blended LTV:CAC: 8.3x ✓✓', style='List Bullet')
doc.add_paragraph('Payback Period: 3.2 months', style='List Bullet')
doc.add_paragraph('Gross Margin: 90% (SaaS standard)', style='List Bullet')

doc.add_page_break()

# COMPETITIVE ANALYSIS
add_heading(doc, 'COMPETITIVE ANALYSIS', 1)

add_heading(doc, 'Why Incumbents Cannot Compete', 2)

# Competitive table
table = doc.add_table(rows=11, cols=5)
table.style = 'Light Grid Accent 1'

comp_headers = ['Factor', 'Coupa', 'Ariba', 'Jaggr', 'Kraftd']
for i, header in enumerate(comp_headers):
    table.rows[0].cells[i].text = header
    set_cell_background(table.rows[0].cells[i], '4472C4')
    table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

comp_data = [
    ('Architecture', 'Enterprise-first', 'Enterprise-first', 'Enterprise-first', 'Messy-data-first'),
    ('Data Prerequisite', 'Clean schemas', 'Clean schemas', 'Clean schemas', 'Works with chaos'),
    ('Implementation', '3-6 months', '6-12 months', '4-8 weeks', '2-4 weeks'),
    ('Cost to Deploy', '$150K-$1M', '$200K-$1M', '$50K-$200K', '$0-5K'),
    ('Monthly Cost', '$20K-$500K+', '$30K-$1M+', '$10K-$100K', '$5-50K'),
    ('Minimum Company Size', '$50M+ revenue', '$100M+ revenue', '$20M+ revenue', '$500K+ revenue'),
    ('SME Market Focus', '❌ No', '❌ No', '❌ No', '✓ Yes'),
    ('Supplier Intelligence', 'One-way', 'One-way', 'One-way', 'Bidirectional'),
    ('Compliance = Insight', '❌ No', '❌ No', '❌ No', '✓ Yes'),
    ('Works Standalone', '❌ (needs ERP)', '❌ (needs ERP)', '⚠ Partial', '✓ Yes'),
]

for i, row_data in enumerate(comp_data, 1):
    for j, cell_value in enumerate(row_data):
        table.rows[i].cells[j].text = cell_value
        if j == 4 and '✓' in cell_value:  # Kraftd strength
            set_cell_background(table.rows[i].cells[j], 'C6E0B4')

doc.add_paragraph()

add_heading(doc, 'Competitive Advantage: Structural, Not Tactical', 3)

adv_reasons = [
    'Why Coupa cannot enter this market:\n• Their sales motion assumes $500K minimum deal\n• Their support model requires ERP integration team\n• Their architecture breaks with messy data\n• Their customers (enterprises) don\'t want them serving suppliers',
    
    'Why Ariba cannot enter this market:\n• Entirely embedded in SAP ecosystem\n• Designed for procurement within ERP, not independent intelligence\n• Would canibalize higher-margin enterprise sales',
    
    'Why new startups struggle:\n• SME sales = high-touch, low-deal-size\n• Requires deep supply chain domain expertise\n• Margins work only at scale\n• Most investors prefer $500K+ deal SaaS',
    
    'Why Kraftd owns this space:\n• Built for this market from inception\n• Architecture supports messy data (not a patch)\n• Economics work (90% margin at scale)\n• First-mover advantage with no follower in sight',
]

for i, reason in enumerate(adv_reasons, 1):
    para = doc.add_paragraph(reason)
    if i == 4:
        para.runs[0].font.bold = True

doc.add_page_break()

# SWOT ANALYSIS
add_heading(doc, 'SWOT ANALYSIS (Revised)', 1)

add_heading(doc, 'STRENGTHS ✓', 2)

strengths = [
    ('Market Positioning', ['First-mover in SME intelligence translation layer', 'Solving structural problem (not incremental)', 'Addressing $2-3T market ignored by incumbents', 'Manifesto-driven mission (attracts founders, employees, investors)']),
    
    ('Product', ['MVP live in production', '95%+ document accuracy (proven)', 'Works with messy reality (not schema-dependent)', '2-4 week implementation (vs 3-12 months competitors)', '$0-5K setup cost (vs $500K-$1M competitors)', 'Standalone or integrated (not ERP-dependent)']),
    
    ('Economics', ['90% gross margin (exceptional for SaaS)', '8.3x blended LTV:CAC (excellent)', '3.2 month payback (capital efficient)', 'Profitable by Year 1 ($560K EBITDA)', 'Scales without proportional cost increase']),
    
    ('Team & Mission', ['Founder proven execution (MVP in 12 months)', 'Clear vision (solves real problem, not incremental)', 'Supply chain domain expertise', 'Fairness/equity narrative attracts top talent']),
]

for strength_category, strength_items in strengths:
    para = doc.add_paragraph(strength_category)
    para.runs[0].font.bold = True
    for item in strength_items:
        doc.add_paragraph(item, style='List Bullet')

add_heading(doc, 'WEAKNESSES ⚠', 2)

weaknesses = [
    ('Market', ['SME sales = high-touch, requires field team', 'Unknown brand (vs $9B Coupa)', 'No customer logos or case studies yet', 'SME sales cycles = 4-8 weeks']),
    
    ('Product', ['Phased feature rollout (full feature set in 12-16 weeks)', 'Advanced analytics not yet live (Phases 2-5)', 'Mobile app not yet available', 'Limited integration ecosystem (early)']),
    
    ('Operations', ['Small team (needs hiring for scale)', 'Limited customer success infrastructure', 'No PR/analyst relations', 'No channel partnership network yet']),
]

for weakness_category, weakness_items in weaknesses:
    para = doc.add_paragraph(weakness_category)
    para.runs[0].font.bold = True
    for item in weakness_items:
        doc.add_paragraph(item, style='List Bullet')

add_heading(doc, 'OPPORTUNITIES 🎯', 2)

opportunities = [
    ('Market Expansion', ['SME segment = $2-3T globally vs $550B initial estimate', 'Geography: Saudi, UAE, broader MENA, APAC', 'Vertical expansion: Agriculture, manufacturing, logistics, construction, trading', 'Adjacent use cases: Manufacturing operations, logistics networks, trading platforms']),
    
    ('Product & Ecosystem', ['Enterprise supply chain transparency (B2B2B model)', 'Supply chain intelligence data/API (B2B data business)', 'Strategic sourcing academy (training, certifications)', 'Supplier network marketplace (matching buyers/sellers)', 'Sustainability intelligence (ESG compliance, carbon tracking)']),
    
    ('Strategic', ['SAP/Oracle integration partnerships (resell through channels)', 'System integrators (Deloitte, EY, Accenture partnerships)', 'Government digitalization programs (Saudi Vision 2030, etc.)', 'Supply chain resilience initiatives (post-pandemic focus)']),
]

for opportunity_category, opportunity_items in opportunities:
    para = doc.add_paragraph(opportunity_category)
    para.runs[0].font.bold = True
    for item in opportunity_items:
        doc.add_paragraph(item, style='List Bullet')

add_heading(doc, 'THREATS 🚨', 2)

threats = [
    ('Competitive', ['Coupa/Ariba could build separate "SME division" (unlikely but possible)', 'Well-funded startup competitors (3-5 funding attempts in space)', 'Microsoft/OpenAI making AI tools ubiquitous (anyone can add chatbots)', 'Low barriers to imitation for basic features']),
    
    ('Market', ['Economic recession reduces SME spending', 'SME consolidation (larger buyers reduce SME count)', 'Enterprise pressure on suppliers to use enterprise systems (instead of Kraftd)', 'Regulatory changes (data localization, privacy)']),
    
    ('Operational', ['Execution risk on product roadmap (Phases 2-5)', 'Churn if early customers dissatisfied', 'Team attrition (key person risk)', 'Fundraising dependency (if don\'t hit metrics)']),
]

for threat_category, threat_items in threats:
    para = doc.add_paragraph(threat_category)
    para.runs[0].font.bold = True
    for item in threat_items:
        doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# FEASIBILITY STUDY
add_heading(doc, 'FEASIBILITY STUDY', 1)

add_heading(doc, 'Technical Feasibility ✓ PROVEN', 2)
tech_items = [
    'Document processing: 95%+ accuracy demonstrated',
    'AI/ML infrastructure: Operational',
    'Multi-tenant architecture: Scalable to 1000+ users',
    'Security/compliance: Enterprise-grade',
    'Cost structure: $37-68/month base (proven scalable)',
]
for item in tech_items:
    doc.add_paragraph(item, style='List Bullet')
doc.add_paragraph()
doc.add_paragraph('Remaining work: Phases 2-5 features (12-16 weeks), Mobile app (8-10 weeks), Enterprise integration ecosystem (ongoing)')

add_heading(doc, 'Market Feasibility ✓ VALIDATED', 2)
market_items = [
    'Target customers exist (SMEs, suppliers, traders)',
    'Problem is real (manifesto validated by real experience)',
    'Willingness to pay proven ($5-50K/month range realistic)',
    'Market timing right (digitalization wave, ESG mandates, supply chain resilience focus)',
    'No competitor in "translation layer" space',
]
for item in market_items:
    doc.add_paragraph(item, style='List Bullet')
doc.add_paragraph()
doc.add_paragraph('Remaining validation: Scale to 500+ customers, Prove unit economics at scale, Build channel partnerships')

add_heading(doc, 'Financial Feasibility ✓ EXCELLENT', 2)
fin_items = [
    '90% gross margin achievable (proven in Year 1)',
    'Positive EBITDA by Month 18 (Year 1 $560K shown)',
    'Payback in 3.2 months (capital efficient)',
    'No high CapEx (cloud-native, SaaS)',
    'Scales with minimal team increase (software leverage)',
]
for item in fin_items:
    doc.add_paragraph(item, style='List Bullet')
doc.add_paragraph()
doc.add_paragraph('Path to profitability proven.')

add_heading(doc, 'Operational Feasibility ⚠ REQUIRES INVESTMENT', 2)
doc.add_paragraph('What\'s needed:')
ops_items = ['Sales team (3-5 people by Month 6)', 'Marketing/demand generation (1-2 people by Month 3)', 'Customer success (2-3 people by Month 6)', 'Operations/finance (1 person by Month 4)', 'Additional engineers (1-2 by Month 6)']
for item in ops_items:
    doc.add_paragraph(item, style='List Bullet')
doc.add_paragraph()
doc.add_paragraph('Cost: $400K Year 1 (sales/marketing) + $200K (product/ops) = $600K')
doc.add_paragraph('Runway: With $2M Series A, 18+ months to profitability')

doc.add_page_break()

# INVESTMENT OPPORTUNITY
add_heading(doc, 'INVESTMENT OPPORTUNITY', 1)

add_heading(doc, 'The Thesis', 2)

para = doc.add_paragraph('Kraftd is building the intelligence translation layer for the $2-3 trillion fragmented supply chain segment.')
para.runs[0].font.bold = True

doc.add_paragraph()
doc.add_paragraph('This market is structurally ignored by incumbents because:')
doc.add_paragraph('1. Their architecture doesn\'t work with messy data', style='List Number')
doc.add_paragraph('2. Their unit economics require $500K+ deals (not $5K-50K)', style='List Number')
doc.add_paragraph('3. Serving SMEs would undermine enterprise margin', style='List Number')
doc.add_paragraph('4. They are committed to enterprise, not suppliers', style='List Number')

doc.add_paragraph()
para = doc.add_paragraph('Kraftd is the only company architecturally and economically built for this space.')
para.runs[0].font.bold = True

add_heading(doc, 'Why Invest Now', 2)

invest_reasons = [
    ('Market Timing', 'Digitalization wave (73% of enterprises currently investing)\nESG compliance mandates driving transparency\nSupply chain resilience post-pandemic\nSME cost pressure creating urgency'),
    
    ('Founder Proven', 'Built MVP to production in 12 months\nClear vision and execution\nDomain expertise in supply chain\nCoachable and mission-driven'),
    
    ('Product Magic', '95%+ accuracy on messy documents\nWorks standalone (no ERP prerequisite)\n2-4 week implementation\n$0-5K setup cost'),
    
    ('Business Model', '90% gross margin\n8.3x blended LTV:CAC\n3.2 month payback\nProfitable by Year 1'),
    
    ('Market Opportunity', '$2-3T addressable (not $550B)\nStructurally ignored by competitors\nMultiple revenue streams\nGeographic expansion paths'),
]

for reason_title, reason_details in invest_reasons:
    para = doc.add_paragraph(reason_title)
    para.runs[0].font.bold = True
    doc.add_paragraph(reason_details)

add_heading(doc, 'Investment Ask', 2)

para = doc.add_paragraph('Series A: $2-3M')
para.runs[0].font.bold = True
para.runs[0].font.size = Pt(12)

doc.add_paragraph()
doc.add_paragraph('Use of Funds:')
funds_items = [
    'Sales/GTM team: $800K (expand from founder to 3-5 reps)',
    'Marketing/demand gen: $400K (build brand, create demand)',
    'Product development: $300K (Phase 2-3 features)',
    'Customer success: $300K (onboarding, retention)',
    'Operations/finance: $200K (build infrastructure)',
]
for item in funds_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph('Total Year 1: $2M')
doc.add_paragraph('Runway: 18+ months to profitability')

doc.add_paragraph()
doc.add_paragraph('Expected Outcomes:')
outcomes = [
    'Year 1: 560 customers, $6.4M revenue, $560K EBITDA',
    'Year 2: 1,200 customers, $19.3M revenue, $8M EBITDA',
    'Year 3: 2,500 customers, $60M revenue, $34M EBITDA',
    'Year 5: 5,000+ customers, $140M revenue, $80M EBITDA',
]
for outcome in outcomes:
    doc.add_paragraph(outcome, style='List Bullet')

add_heading(doc, 'Return Scenarios', 2)

returns = [
    ('Base Case (Most Likely)', 'Exit Year: 5\nExit Valuation: $2-3B (strategic acquisition)\nInvestor Return: 25-35x on $2M investment'),
    
    ('Upside Case (Strong Execution)', 'Exit Year: 4-5\nExit Valuation: $3-5B\nInvestor Return: 50x+'),
    
    ('IPO Case (Market Leadership)', 'Exit Year: 6-7\nExit Valuation: $5-10B (SaaS public multiples)\nInvestor Return: 100x+'),
]

for case_name, case_details in returns:
    para = doc.add_paragraph(case_name)
    para.runs[0].font.bold = True
    doc.add_paragraph(case_details)

doc.add_page_break()

# EXECUTIVE SUMMARY
add_heading(doc, 'EXECUTIVE SUMMARY FOR INVESTORS', 1)

add_heading(doc, 'The Problem', 2)
doc.add_paragraph('SMEs in supply chains operate without the intelligence that large enterprises take for granted.')
doc.add_paragraph()
doc.add_paragraph('SABIC has structured data, APIs, dashboards, and predictive analytics.')
doc.add_paragraph('An SME supplier doing identical supply chain work has spreadsheets, emails, and manual tracking.')
doc.add_paragraph()
para = doc.add_paragraph('This isn\'t inefficiency. This is structural inequality in information access.')
para.runs[0].font.bold = True

add_heading(doc, 'The Opportunity', 2)
doc.add_paragraph('$2-3 trillion in fragmented supply chains where intelligence has never reached.')
doc.add_paragraph()
doc.add_paragraph('SMEs, suppliers, traders, logistics providers, manufacturers — all operating without visibility into price trends, risk signals, supplier reliability, or competitive positioning.')
doc.add_paragraph()
doc.add_paragraph('Incumbents (Coupa, SAP, Oracle) cannot serve this market:')
doc.add_paragraph('• Their architecture requires clean data (SMEs have chaos)', style='List Bullet')
doc.add_paragraph('• Their unit economics require $500K+ deals (SMEs have $5-50K budgets)', style='List Bullet')
doc.add_paragraph('• Their mission is enterprise (suppliers are cost centers)', style='List Bullet')

add_heading(doc, 'The Solution', 2)
doc.add_paragraph('Kraftd is a translation layer between messy reality and intelligence.')
doc.add_paragraph()
doc.add_paragraph('We work before APIs, before schemas, before enterprise infrastructure — where 80% of actual supply chain work happens.')
doc.add_paragraph()
doc.add_paragraph('• Take messy documents (PDFs, invoices, contracts, customs forms)', style='List Bullet')
doc.add_paragraph('• Extract structured intelligence (95%+ accuracy)', style='List Bullet')
doc.add_paragraph('• Provide SME-grade decision support', style='List Bullet')
doc.add_paragraph('• No ERP prerequisite, no IT team required, 2-4 week implementation', style='List Bullet')

add_heading(doc, 'The Business Model', 2)
doc.add_paragraph('3 revenue streams:')
doc.add_paragraph('SME Intelligence (70%): $499-$9,999/month tiered subscriptions', style='List Number')
doc.add_paragraph('Enterprise Transparency (20%): Enterprises pay for supplier visibility', style='List Number')
doc.add_paragraph('Supply Chain Data (10%): Intelligence products, price indices, risk signals', style='List Number')

add_heading(doc, 'The Numbers', 2)

final_table = doc.add_table(rows=6, cols=5)
final_table.style = 'Light Grid Accent 1'

final_headers = ['Year', 'Customers', 'ARR', 'Revenue', 'EBITDA']
for i, header in enumerate(final_headers):
    final_table.rows[0].cells[i].text = header
    set_cell_background(final_table.rows[0].cells[i], 'D3D3D3')

final_data = [
    ('1', '560', '$6M', '$6.4M', '$560K'),
    ('2', '1,200', '$18.5M', '$19.3M', '$8M'),
    ('3', '2,500', '$58.5M', '$60M', '$34M'),
    ('5', '5,000+', '$120M+', '$140M+', '$80M+'),
]

for i, row_data in enumerate(final_data, 1):
    for j, value in enumerate(row_data):
        final_table.rows[i].cells[j].text = value

doc.add_paragraph()
doc.add_paragraph('Unit Economics:')
doc.add_paragraph('LTV:CAC: 8.3x (blended)', style='List Bullet')
doc.add_paragraph('Payback: 3.2 months', style='List Bullet')
doc.add_paragraph('Gross Margin: 90%', style='List Bullet')
doc.add_paragraph('Profitability: Month 18', style='List Bullet')

add_heading(doc, 'Why Kraftd Wins', 2)

wins = [
    'First-mover in translation layer space (no competitors)',
    'Structural moat (incumbents can\'t compete here)',
    'Proven founder (MVP in production)',
    'Exceptional economics (profitable by Year 1)',
    'Massive TAM ($2-3T, not $550B)',
    'Mission-driven (solves real inequality, not incremental)',
]
for win in wins:
    doc.add_paragraph(win, style='List Number')

add_heading(doc, 'The Investment', 2)

para = doc.add_paragraph('Series A: $2-3M')
para.runs[0].font.bold = True
doc.add_paragraph()
doc.add_paragraph('Expected Return: 25-35x in base case, 50-100x+ in upside cases')
doc.add_paragraph()
doc.add_paragraph('Exit Timeline: 5-7 years')

doc.add_page_break()

# CLOSING
add_heading(doc, 'CLOSING', 1)

doc.add_paragraph('Kraftd is not building "better procurement software."')
doc.add_paragraph()
para = doc.add_paragraph('Kraftd is building intelligence equity in supply chains.')
para.runs[0].font.bold = True
para.runs[0].font.size = Pt(14)

doc.add_paragraph()
doc.add_paragraph('We are solving a structural problem: the systematic exclusion of SMEs and suppliers from the intelligence that enterprises have.')
doc.add_paragraph()
doc.add_paragraph('This is a $2-3 trillion opportunity in fragmented supply chains.')
doc.add_paragraph()
doc.add_paragraph('It is ignored by every incumbent because their architecture, economics, and mission make this market invisible to them.')
doc.add_paragraph()
doc.add_paragraph('We are the only company built for it.')
doc.add_paragraph()
para = doc.add_paragraph('That is the investment thesis.')
para.runs[0].font.bold = True

# Save document
output_path = r'c:\Users\1R6\OneDrive\Project Catalyst\KraftdIntel\KRAFTD_Investment_Analysis_Manifesto_Centered.docx'
doc.save(output_path)
print(f'✓ Word document created: {output_path}')
print(f'✓ Document includes: Manifesto, Market Analysis, Business Model, Financial Projections, Investment Thesis')
print(f'✓ Ready for investor presentations and distribution')
