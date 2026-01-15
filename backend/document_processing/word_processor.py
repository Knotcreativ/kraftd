from docx import Document
from typing import List, Dict, Any
from .base_processor import BaseProcessor

class WordProcessor(BaseProcessor):
    """Process Word (.docx) documents and extract text, tables, and structured data."""
    
    def parse(self) -> Dict[str, Any]:
        """Parse Word document and return structured data."""
        try:
            doc = Document(self.file_path)
            result = {
                "paragraphs_count": len(doc.paragraphs),
                "tables_count": len(doc.tables),
                "text": "",
                "tables": [],
                "paragraphs": []
            }
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    result["text"] += para.text + "\n"
                    result["paragraphs"].append({
                        "text": para.text,
                        "style": para.style.name
                    })
            
            # Extract tables
            for table_idx, table in enumerate(doc.tables):
                table_data = []
                for row_idx, row in enumerate(table.rows):
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                result["tables"].append({
                    "table_number": table_idx + 1,
                    "data": table_data
                })
            
            return result
        except Exception as e:
            return {"error": f"Failed to parse Word document: {str(e)}"}
    
    def extract_tables(self) -> List[List[Dict]]:
        """Extract all tables from the Word document."""
        tables = []
        try:
            doc = Document(self.file_path)
            for table_idx, table in enumerate(doc.tables, 1):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables.append({
                    "table_number": table_idx,
                    "data": table_data
                })
        except Exception as e:
            return [{"error": f"Failed to extract tables: {str(e)}"}]
        return tables
    
    def extract_text(self) -> str:
        """Extract all text from the Word document."""
        text = ""
        try:
            doc = Document(self.file_path)
            for para in doc.paragraphs:
                text += para.text + "\n"
        except Exception as e:
            text = f"Error extracting text: {str(e)}"
        return text
