"""
Document Ingestion & Enhancement System

Loads procurement documents from device storage to enhance ML and AI models.
Supports: Local files, OneDrive, Azure Blob Storage, Network shares.
Formats: PDF, Images, Excel, CSV, JSON, Word documents.
"""

import os
import json
import logging
from pathlib import Path
from typing import Dict, Any, List, Optional, Tuple
from dataclasses import dataclass, asdict
from datetime import datetime
import asyncio
from concurrent.futures import ThreadPoolExecutor

logger = logging.getLogger(__name__)


@dataclass
class ProcurementDocument:
    """Structured representation of a procurement document"""
    file_path: str
    file_name: str
    file_type: str  # pdf, image, excel, csv, json, etc.
    document_type: str  # quotation, po, rfq, contract, invoice
    file_size: int
    extracted_data: Dict[str, Any]
    supplier_name: Optional[str]
    total_amount: Optional[float]
    date_extracted: str
    raw_text: Optional[str]
    metadata: Dict[str, Any]


class DocumentIngestionService:
    """
    Manages document ingestion from various storage sources.
    Prepares data for ML training and AI context enhancement.
    """
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
        self.documents: List[ProcurementDocument] = []
        self.storage_path = None
        self.executor = ThreadPoolExecutor(max_workers=5)
        
    async def connect_to_storage(
        self,
        storage_path: str,
        storage_type: str = "local"
    ) -> bool:
        """
        Connect to document storage.
        
        Args:
            storage_path: Path to storage (local path, OneDrive path, blob URL, etc.)
            storage_type: "local", "onedrive", "azure_blob", "network_share"
            
        Returns:
            bool: True if connection successful
        """
        try:
            if storage_type == "local":
                path = Path(storage_path)
                if not path.exists():
                    self.logger.error(f"Path does not exist: {storage_path}")
                    return False
                if not path.is_dir():
                    self.logger.error(f"Path is not a directory: {storage_path}")
                    return False
                self.storage_path = path
                self.logger.info(f"✓ Connected to local storage: {storage_path}")
                return True
                
            elif storage_type == "onedrive":
                # OneDrive path (part of Windows filesystem)
                path = Path(storage_path)
                if not path.exists():
                    self.logger.error(f"OneDrive path does not exist: {storage_path}")
                    return False
                self.storage_path = path
                self.logger.info(f"✓ Connected to OneDrive: {storage_path}")
                return True
                
            elif storage_type == "azure_blob":
                self.logger.info("Azure Blob Storage support requires azure-storage-blob SDK")
                # TODO: Implement Azure Blob Storage connection
                return False
                
            else:
                self.logger.error(f"Unsupported storage type: {storage_type}")
                return False
                
        except Exception as e:
            self.logger.error(f"Failed to connect to storage: {e}")
            return False
    
    async def scan_documents(
        self,
        file_extensions: Optional[List[str]] = None,
        recursive: bool = True
    ) -> int:
        """
        Scan storage for procurement documents.
        
        Args:
            file_extensions: List of extensions to scan for (e.g., ['.pdf', '.xlsx'])
            recursive: Scan subdirectories
            
        Returns:
            Number of documents found
        """
        if not self.storage_path:
            self.logger.error("Storage not connected. Call connect_to_storage() first.")
            return 0
        
        if file_extensions is None:
            file_extensions = [
                '.pdf', '.png', '.jpg', '.jpeg', '.tiff', '.bmp',
                '.xlsx', '.xls', '.csv', '.json', '.docx', '.doc'
            ]
        
        self.logger.info(f"Scanning for documents in: {self.storage_path}")
        self.logger.info(f"Extensions: {file_extensions}")
        
        try:
            pattern = "**/*" if recursive else "*"
            files = list(self.storage_path.glob(pattern))
            
            filtered_files = [
                f for f in files
                if f.is_file() and f.suffix.lower() in file_extensions
            ]
            
            self.logger.info(f"Found {len(filtered_files)} documents")
            
            # Process files asynchronously
            tasks = [
                self._load_document(file_path)
                for file_path in filtered_files[:100]  # Limit to first 100 for now
            ]
            
            loaded_docs = await asyncio.gather(*tasks, return_exceptions=True)
            self.documents = [d for d in loaded_docs if isinstance(d, ProcurementDocument)]
            
            self.logger.info(f"Successfully loaded {len(self.documents)} documents")
            return len(self.documents)
            
        except Exception as e:
            self.logger.error(f"Error scanning documents: {e}", exc_info=True)
            return 0
    
    async def _load_document(self, file_path: Path) -> Optional[ProcurementDocument]:
        """Load and parse a single document."""
        try:
            file_type = file_path.suffix.lower().lstrip('.')
            
            # Detect document type from filename
            file_name_lower = file_path.stem.lower()
            if any(x in file_name_lower for x in ['quote', 'quotation', 'qtn']):
                doc_type = 'quotation'
            elif any(x in file_name_lower for x in ['po', 'purchase order']):
                doc_type = 'po'
            elif any(x in file_name_lower for x in ['rfq', 'request']):
                doc_type = 'rfq'
            elif any(x in file_name_lower for x in ['contract', 'agreement']):
                doc_type = 'contract'
            elif any(x in file_name_lower for x in ['invoice', 'bill']):
                doc_type = 'invoice'
            else:
                doc_type = 'unknown'
            
            # Extract text from different formats
            extracted_data = {}
            raw_text = None
            
            if file_type == 'pdf':
                raw_text, extracted_data = await self._extract_from_pdf(file_path)
            elif file_type in ['png', 'jpg', 'jpeg', 'tiff', 'bmp']:
                raw_text, extracted_data = await self._extract_from_image(file_path)
            elif file_type in ['xlsx', 'xls']:
                extracted_data = await self._extract_from_excel(file_path)
            elif file_type == 'csv':
                extracted_data = await self._extract_from_csv(file_path)
            elif file_type == 'json':
                extracted_data = await self._extract_from_json(file_path)
            elif file_type in ['docx', 'doc']:
                raw_text, extracted_data = await self._extract_from_word(file_path)
            
            # Try to extract key fields
            supplier_name = extracted_data.get('supplier_name')
            total_amount = extracted_data.get('total_amount')
            
            doc = ProcurementDocument(
                file_path=str(file_path),
                file_name=file_path.name,
                file_type=file_type,
                document_type=doc_type,
                file_size=file_path.stat().st_size,
                extracted_data=extracted_data,
                supplier_name=supplier_name,
                total_amount=total_amount,
                date_extracted=datetime.now().isoformat(),
                raw_text=raw_text,
                metadata={
                    'detected_type': doc_type,
                    'text_length': len(raw_text) if raw_text else 0,
                    'fields_extracted': len(extracted_data),
                    'creation_date': datetime.fromtimestamp(file_path.stat().st_ctime).isoformat()
                }
            )
            
            self.logger.debug(f"Loaded: {file_path.name} ({doc_type})")
            return doc
            
        except Exception as e:
            self.logger.warning(f"Failed to load document {file_path}: {e}")
            return None
    
    async def _extract_from_pdf(self, file_path: Path) -> Tuple[str, Dict]:
        """Extract text and data from PDF."""
        try:
            import PyPDF2
            text = ""
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text += page.extract_text()
            
            # Basic field extraction from text
            extracted = self._extract_fields_from_text(text)
            return text, extracted
        except ImportError:
            self.logger.warning("PyPDF2 not installed. Install with: pip install PyPDF2")
            return "", {}
        except Exception as e:
            self.logger.warning(f"Error extracting from PDF: {e}")
            return "", {}
    
    async def _extract_from_image(self, file_path: Path) -> Tuple[str, Dict]:
        """Extract text from image using OCR."""
        try:
            import pytesseract
            from PIL import Image
            
            image = Image.open(file_path)
            text = pytesseract.image_to_string(image)
            extracted = self._extract_fields_from_text(text)
            return text, extracted
        except ImportError:
            self.logger.warning("pytesseract or PIL not installed. Install with: pip install pytesseract pillow")
            return "", {}
        except Exception as e:
            self.logger.warning(f"Error extracting from image: {e}")
            return "", {}
    
    async def _extract_from_excel(self, file_path: Path) -> Dict:
        """Extract structured data from Excel."""
        try:
            import pandas as pd
            df = pd.read_excel(file_path)
            return df.to_dict('records')
        except ImportError:
            self.logger.warning("pandas not installed. Install with: pip install openpyxl pandas")
            return {}
        except Exception as e:
            self.logger.warning(f"Error extracting from Excel: {e}")
            return {}
    
    async def _extract_from_csv(self, file_path: Path) -> Dict:
        """Extract structured data from CSV."""
        try:
            import pandas as pd
            df = pd.read_csv(file_path)
            return df.to_dict('records')
        except ImportError:
            self.logger.warning("pandas not installed")
            return {}
        except Exception as e:
            self.logger.warning(f"Error extracting from CSV: {e}")
            return {}
    
    async def _extract_from_json(self, file_path: Path) -> Dict:
        """Extract data from JSON."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return json.load(f)
        except Exception as e:
            self.logger.warning(f"Error extracting from JSON: {e}")
            return {}
    
    async def _extract_from_word(self, file_path: Path) -> Tuple[str, Dict]:
        """Extract text from Word document."""
        try:
            from docx import Document
            doc = Document(file_path)
            text = "\n".join([p.text for p in doc.paragraphs])
            extracted = self._extract_fields_from_text(text)
            return text, extracted
        except ImportError:
            self.logger.warning("python-docx not installed. Install with: pip install python-docx")
            return "", {}
        except Exception as e:
            self.logger.warning(f"Error extracting from Word: {e}")
            return "", {}
    
    def _extract_fields_from_text(self, text: str) -> Dict[str, Any]:
        """Extract common procurement fields from text using regex."""
        import re
        
        extracted = {}
        
        # Supplier name (look for patterns like "From:", "Supplier:", "Company:")
        supplier_match = re.search(
            r'(?:From|Supplier|Company|Vendor)[\s:]*([A-Za-z\s&.,()-]+?)(?:\n|$)',
            text, re.IGNORECASE
        )
        if supplier_match:
            extracted['supplier_name'] = supplier_match.group(1).strip()
        
        # Amount (look for currency amounts)
        amount_match = re.search(
            r'(?:Total|Amount|Price|Cost)[\s:]*\$?([\d,]+\.?\d*)',
            text, re.IGNORECASE
        )
        if amount_match:
            try:
                extracted['total_amount'] = float(amount_match.group(1).replace(',', ''))
            except:
                pass
        
        # Date
        date_match = re.search(
            r'(?:Date|Invoice Date)[\s:]*(\d{1,2}[-/]\d{1,2}[-/]\d{2,4})',
            text, re.IGNORECASE
        )
        if date_match:
            extracted['date'] = date_match.group(1)
        
        # Payment terms
        terms_match = re.search(
            r'(?:Terms|Payment)[\s:]*([A-Za-z\s0-9]+?)(?:\n|$)',
            text, re.IGNORECASE
        )
        if terms_match:
            extracted['payment_terms'] = terms_match.group(1).strip()
        
        return extracted
    
    def get_documents(self) -> List[ProcurementDocument]:
        """Get all loaded documents."""
        return self.documents
    
    def get_documents_by_type(self, doc_type: str) -> List[ProcurementDocument]:
        """Get documents of specific type."""
        return [d for d in self.documents if d.document_type == doc_type]
    
    def get_statistics(self) -> Dict[str, Any]:
        """Get statistics about loaded documents."""
        if not self.documents:
            return {"total_documents": 0}
        
        return {
            "total_documents": len(self.documents),
            "by_type": {
                doc_type: len(self.get_documents_by_type(doc_type))
                for doc_type in set(d.document_type for d in self.documents)
            },
            "by_file_type": {
                file_type: len([d for d in self.documents if d.file_type == file_type])
                for file_type in set(d.file_type for d in self.documents)
            },
            "total_size_mb": sum(d.file_size for d in self.documents) / (1024 * 1024),
            "with_supplier_names": len([d for d in self.documents if d.supplier_name]),
            "with_amounts": len([d for d in self.documents if d.total_amount]),
            "total_amount_usd": sum(d.total_amount for d in self.documents if d.total_amount)
        }
    
    def export_for_ml_training(self, output_path: str) -> bool:
        """Export documents as training dataset for ML models."""
        try:
            training_data = []
            
            for doc in self.documents:
                training_record = {
                    "file_name": doc.file_name,
                    "document_type": doc.document_type,
                    "supplier_name": doc.supplier_name,
                    "total_amount": doc.total_amount,
                    "extracted_fields": doc.extracted_data,
                    "file_type": doc.file_type,
                    "file_size": doc.file_size,
                    "metadata": doc.metadata
                }
                training_data.append(training_record)
            
            # Save as JSON lines (one record per line for easy streaming)
            output_file = Path(output_path) / "training_data.jsonl"
            output_file.parent.mkdir(parents=True, exist_ok=True)
            
            with open(output_file, 'w', encoding='utf-8') as f:
                for record in training_data:
                    f.write(json.dumps(record) + '\n')
            
            self.logger.info(f"✓ Exported {len(training_data)} records to {output_file}")
            return True
            
        except Exception as e:
            self.logger.error(f"Error exporting training data: {e}")
            return False
    
    def export_for_ai_context(self, output_path: str, sample_size: int = 10) -> bool:
        """Export sample documents as context examples for gpt-4o."""
        try:
            context_examples = []
            
            # Sample documents from each type
            doc_types = set(d.document_type for d in self.documents)
            
            for doc_type in doc_types:
                docs_of_type = self.get_documents_by_type(doc_type)
                samples = docs_of_type[:sample_size]
                
                for doc in samples:
                    example = {
                        "type": doc_type,
                        "supplier": doc.supplier_name,
                        "amount": doc.total_amount,
                        "extracted_data": doc.extracted_data,
                        "text_sample": (doc.raw_text or "")[:500]  # First 500 chars
                    }
                    context_examples.append(example)
            
            output_file = Path(output_path) / "ai_context_examples.json"
            output_file.parent.mkdir(parents=True, exist_ok=True)
            
            with open(output_file, 'w', encoding='utf-8') as f:
                json.dump(context_examples, f, indent=2)
            
            self.logger.info(f"✓ Exported {len(context_examples)} AI context examples to {output_file}")
            return True
            
        except Exception as e:
            self.logger.error(f"Error exporting AI context: {e}")
            return False


# Global instance
document_ingestion = DocumentIngestionService()
